import os
import logging
from docx import Document

# Logger
logger = logging.getLogger(__name__)


def generate_document(template_path: str, output_path: str, data: dict, user_id: int):
    logger.info(
        "generate_document START template_path=%s output_path=%s user_id=%s",
        template_path,
        output_path,
        user_id
    )

    try:
        # --- Verificar template ---
        logger.info(
            "Checking template exists=%s path=%s",
            os.path.exists(template_path),
            template_path
        )

        if not os.path.exists(template_path):
            logger.error("Template DOCX not found")
            raise FileNotFoundError(template_path)

        # --- Abrir DOCX ---
        logger.info("Opening DOCX template")
        doc = Document(template_path)
        logger.info("DOCX template opened successfully")

        # --- Reemplazar texto en párrafos ---
        replaced_count = 0

        for p in doc.paragraphs:
            original_text = p.text
            for key, value in data.items():
                placeholder = f"${{{key}}}"
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))
            if p.text != original_text:
                replaced_count += 1

        # --- Reemplazar texto en tablas ---
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        original_text = p.text
                        for key, value in data.items():
                            placeholder = f"${{{key}}}"
                            if placeholder in p.text:
                                p.text = p.text.replace(placeholder, str(value))
                        if p.text != original_text:
                            replaced_count += 1

        logger.info("Text replacement completed replaced_count=%s", replaced_count)

        # --- Guardar DOCX ---
        output_dir = os.path.dirname(output_path)
        logger.info("Ensuring output_dir exists=%s", output_dir)
        os.makedirs(output_dir, exist_ok=True)

        logger.info("Saving DOCX to output_path=%s", output_path)
        doc.save(output_path)

        logger.info(
            "DOCX saved exists=%s size=%s",
            os.path.exists(output_path),
            os.path.getsize(output_path) if os.path.exists(output_path) else None
        )

        if not os.path.exists(output_path):
            logger.error("DOCX save failed")
            raise RuntimeError("DOCX was not saved")

        logger.info("generate_document SUCCESS")

    except Exception:
        logger.exception("generate_document FAILED")
        raise
