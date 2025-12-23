from docx import Document
import os


def generate_document(template_file, output_file, data, user_id):
    """
    template_file: ruta completa al .docx
    output_file: ruta completa de salida
    data: dict con placeholders
    user_id: id del usuario (para auditoría futura)
    """

    # Abrir directamente la plantilla (ya fue validada antes)
    doc = Document(template_file)

    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"${{{key}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"${{{key}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

    # Crear carpeta destino si no existe
    os.makedirs(os.path.dirname(output_file), exist_ok=True)

    # Guardar documento final
    doc.save(output_file)

    return output_file
