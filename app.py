"""
app.py — Zeon (versión simple, sin AWS) listo para Render

Qué hace esta app (Flask):
- Sirve frontend estático (opcional) desde /static
- Expone endpoints HTTP para:
  1) Listar tipos de plantillas disponibles
  2) Generar un .docx desde una plantilla reemplazando placeholders ${...}
  3) Descargar el documento generado
  4) (Opcional) Enviar el .docx por email vía SMTP (Gmail u otro)

Cómo desplegar en Render:
- requirements.txt mínimo:
    Flask
    python-docx
    gunicorn
- Start command:
    gunicorn app:app
- Variables de entorno (opcionales):
    FLASK_ENV=production
    APP_BASE_URL=https://<tu-servicio>.onrender.com   (para armar links absolutos)
    TEMPLATE_ROOT=plantillas
    OUTPUT_ROOT=outputs

  Email (si querés que envíe):
    SMTP_ENABLED=1
    SMTP_HOST=smtp.gmail.com
    SMTP_PORT=465
    SMTP_USER=tuemail@gmail.com
    SMTP_PASS=tu_app_password
    SMTP_FROM="Zeon <tuemail@gmail.com>"   (opcional)
    SMTP_TLS=0  (si usás 465 SSL)
    SMTP_SSL=1  (recomendado Gmail)
    SMTP_DEBUG=0
"""

from __future__ import annotations

import os
import re
import uuid
import json
import time
import shutil
import logging
import mimetypes
import traceback
from dotenv import load_dotenv
load_dotenv()
from dataclasses import dataclass
from typing import Dict, Any, Optional, List, Tuple
from flask import send_from_directory
from routes.admin_templates_routes import bp as admin_templates_bp
from core.auth import create_user, get_user_by_email

from flask import (
    Flask,
    request,
    jsonify,
    send_file,
    abort,
    Response,
)
from docx import Document # type: ignore
from core.auth import ensure_schema, bootstrap_admin_if_needed
from routes.auth_routes import bp as auth_bp
from routes.client_routes import client_bp
from routes.superadmin_routes import bp as superadmin_bp


# -----------------------------------------------------------------------------
# Configuración general
# -----------------------------------------------------------------------------



def _env_bool(name: str, default: bool = False) -> bool:
    v = os.getenv(name, "").strip().lower()
    if not v:
        return default
    return v in ("1", "true", "yes", "y", "on")


def _env_int(name: str, default: int) -> int:
    v = os.getenv(name, "").strip()
    if not v:
        return default
    try:
        return int(v)
    except ValueError:
        return default


APP_BASE_URL = os.getenv("APP_BASE_URL", "").strip().rstrip("/")
TEMPLATE_ROOT = os.getenv("TEMPLATE_ROOT", "plantillas").strip()
OUTPUT_ROOT = os.getenv("OUTPUT_ROOT", "outputs").strip()

MAX_OUTPUT_FILES = _env_int("MAX_OUTPUT_FILES", 5000)  # límite lógico (no hard)
MAX_OUTPUT_AGE_SECONDS = _env_int("MAX_OUTPUT_AGE_SECONDS", 7 * 24 * 3600)  # 7 días

# Email (opcional)
SMTP_ENABLED = _env_bool("SMTP_ENABLED", False)
SMTP_HOST = os.getenv("SMTP_HOST", "smtp.gmail.com").strip()
SMTP_PORT = _env_int("SMTP_PORT", 465)
SMTP_USER = os.getenv("SMTP_USER", "").strip()
SMTP_PASS = os.getenv("SMTP_PASS", "").strip()
SMTP_FROM = os.getenv("SMTP_FROM", "").strip()  # si vacío, usa SMTP_USER
SMTP_SSL = _env_bool("SMTP_SSL", True)          # Gmail: True con 465
SMTP_TLS = _env_bool("SMTP_TLS", False)         # True si usás STARTTLS
SMTP_DEBUG = _env_bool("SMTP_DEBUG", False)

# Seguridad / CORS simple (si necesitás llamar desde otro dominio)
CORS_ENABLED = _env_bool("CORS_ENABLED", True)
CORS_ALLOW_ORIGIN = os.getenv("CORS_ALLOW_ORIGIN", "*").strip()

# Logging
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").strip().upper()


# -----------------------------------------------------------------------------
# Logging setup
# -----------------------------------------------------------------------------

logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("zeon")


# -----------------------------------------------------------------------------
# Flask app
# -----------------------------------------------------------------------------

app = Flask(
    __name__,
    static_folder="static",
    static_url_path=""
)


def ensure_superadmin():
    email = "root@demo.com"
    password = "Root123!"
    role = "superadmin"

    try:
        if not get_user_by_email(email):
            create_user(email, password, role=role)
            print("Superadmin creado automáticamente")
    except Exception as e:
        print("Superadmin ya existe o error:", e)



# -----------------------------------------------------------------------------
# Helpers de archivos / paths
# -----------------------------------------------------------------------------

FILENAME_SAFE_RE = re.compile(r"[^a-zA-Z0-9._-]+")

ensure_schema()
bootstrap_admin_if_needed()
ensure_superadmin()
app.register_blueprint(auth_bp)
app.register_blueprint(client_bp)
app.register_blueprint(superadmin_bp)
app.register_blueprint(admin_templates_bp)


def sanitize_filename(name: str, max_len: int = 200) -> str:
    """
    Normaliza nombres de archivo para evitar caracteres raros.
    Mantiene: letras, números, punto, guión, guión bajo.
    """
    name = (name or "").strip()
    name = name.replace(" ", "_")
    name = FILENAME_SAFE_RE.sub("_", name)
    name = re.sub(r"[_]{2,}", "_", name)
    name = name.strip("._-")
    if not name:
        name = "file"
    return name[:max_len]


def ensure_dirs() -> None:
    os.makedirs(TEMPLATE_ROOT, exist_ok=True)
    os.makedirs(OUTPUT_ROOT, exist_ok=True)


def safe_join(base: str, *paths: str) -> str:
    """
    Join seguro para evitar path traversal.
    """
    base_abs = os.path.abspath(base)
    joined = os.path.abspath(os.path.join(base_abs, *paths))
    if not joined.startswith(base_abs + os.sep) and joined != base_abs:
        raise ValueError("Ruta inválida (path traversal detectado).")
    return joined


def list_template_types() -> List[str]:
    """
    Lista carpetas dentro de TEMPLATE_ROOT.
    Cada carpeta representa un tipo de plantilla.
    """
    ensure_dirs()
    types: List[str] = []
    for name in os.listdir(TEMPLATE_ROOT):
        p = os.path.join(TEMPLATE_ROOT, name)
        if os.path.isdir(p) and not name.startswith("."):
            types.append(name)
    types.sort()
    return types


def find_template_docx(template_type: str) -> str:
    """
    Busca un .docx dentro de plantillas/<template_type>/.
    - Si existe 'plantilla.docx' o 'template.docx' lo prioriza.
    - Si no, toma el primer .docx que encuentre.
    """
    ensure_dirs()
    template_type = sanitize_filename(template_type)
    folder = safe_join(TEMPLATE_ROOT, template_type)
    if not os.path.isdir(folder):
        raise FileNotFoundError(f"No existe la carpeta de plantilla: {template_type}")

    preferred = ["plantilla.docx", "template.docx", "plantilla_solicitud.docx"]
    for fn in preferred:
        candidate = os.path.join(folder, fn)
        if os.path.isfile(candidate):
            return candidate

    docxs = [f for f in os.listdir(folder) if f.lower().endswith(".docx")]
    docxs.sort()
    if not docxs:
        raise FileNotFoundError(f"No se encontró ningún .docx en: {folder}")
    return os.path.join(folder, docxs[0])


def build_output_filename(prefix: str = "documento", ext: str = ".docx") -> str:
    """
    Genera un nombre único y razonablemente legible.
    """
    prefix = sanitize_filename(prefix)[:60]
    unique = uuid.uuid4().hex
    return f"{prefix}_{unique}{ext}"


def output_path_for(filename: str) -> str:
    """
    Resuelve un archivo dentro de OUTPUT_ROOT de forma segura.
    """
    filename = sanitize_filename(filename)
    return safe_join(OUTPUT_ROOT, filename)


def cleanup_outputs() -> None:
    """
    Limpieza best-effort de archivos viejos en OUTPUT_ROOT.
    Render tiene disco efímero; igual conviene mantener orden.
    """
    try:
        ensure_dirs()
        now = time.time()
        files = []
        for fn in os.listdir(OUTPUT_ROOT):
            path = os.path.join(OUTPUT_ROOT, fn)
            if os.path.isfile(path):
                try:
                    st = os.stat(path)
                    files.append((path, st.st_mtime))
                except Exception:
                    continue

        # borrar viejos
        for path, mtime in files:
            if (now - mtime) > MAX_OUTPUT_AGE_SECONDS:
                try:
                    os.remove(path)
                except Exception:
                    pass

        # si hay demasiados, borrar los más viejos
        files = [(p, m) for (p, m) in files if os.path.exists(p)]
        if len(files) > MAX_OUTPUT_FILES:
            files.sort(key=lambda x: x[1])  # más viejos primero
            to_delete = files[: len(files) - MAX_OUTPUT_FILES]
            for path, _ in to_delete:
                try:
                    os.remove(path)
                except Exception:
                    pass
    except Exception:
        logger.debug("cleanup_outputs() falló: %s", traceback.format_exc())


# -----------------------------------------------------------------------------
# Reemplazo robusto en Word (python-docx)
# -----------------------------------------------------------------------------

PLACEHOLDER_RE = re.compile(r"\$\{([A-Za-z0-9_.-]+)\}")


def normalize_replacements(user_data: Dict[str, Any]) -> Dict[str, str]:
    """
    Convierte un dict de datos a placeholders ${KEY} -> 'valor' (str).
    - Acepta claves con o sin ${}
    - Convierte None -> ""
    - Convierte tipos no str a str
    """
    replacements: Dict[str, str] = {}

    for k, v in (user_data or {}).items():
        if k is None:
            continue
        key = str(k).strip()
        if not key:
            continue

        # permitimos que el usuario pase "NOMBRE" o "${NOMBRE}"
        if key.startswith("${") and key.endswith("}"):
            placeholder = key
        else:
            placeholder = "${" + key + "}"

        if v is None:
            replacements[placeholder] = ""
        elif isinstance(v, (dict, list)):
            # por si mandan objetos, los serializamos bonito
            replacements[placeholder] = json.dumps(v, ensure_ascii=False)
        else:
            replacements[placeholder] = str(v)

    return replacements


def replace_placeholders_in_text(text: str, replacements: Dict[str, str]) -> Tuple[str, List[str]]:
    """
    Reemplaza placeholders dentro de un string y devuelve:
    - texto resultante
    - lista de placeholders que fueron reemplazados
    """
    replaced_keys: List[str] = []

    def _sub(m: re.Match) -> str:
        key_inside = m.group(1)
        placeholder = "${" + key_inside + "}"
        if placeholder in replacements:
            replaced_keys.append(placeholder)
            return replacements[placeholder]
        # si no está, lo dejamos igual (para detectar faltantes)
        return m.group(0)

    out = PLACEHOLDER_RE.sub(_sub, text)
    return out, replaced_keys


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> List[str]:
    """
    Reemplaza placeholders en un párrafo considerando que Word puede partir
    el texto en múltiples runs.

    Estrategia:
    - Concatenamos texto total del párrafo (runs)
    - Hacemos replace sobre el texto plano
    - Si cambia, reescribimos el párrafo:
        - ponemos el texto completo en el primer run
        - vaciamos el resto
      Esto puede "aplanar" estilos dentro del párrafo si estaban mezclados.
      En la práctica suele ser aceptable para plantillas de contratos.
    """
    if not paragraph.runs:
        return []

    original = "".join(run.text for run in paragraph.runs)
    new_text, replaced = replace_placeholders_in_text(original, replacements)

    if new_text == original:
        return []

    # Reescribimos con el mínimo impacto posible
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""

    return replaced


def _replace_in_table(table, replacements: Dict[str, str]) -> List[str]:
    replaced_all: List[str] = []
    for row in table.rows:
        for cell in row.cells:
            # Párrafos dentro de la celda
            for p in cell.paragraphs:
                replaced_all.extend(_replace_in_paragraph(p, replacements))
            # Tablas anidadas dentro de la celda
            for t in cell.tables:
                replaced_all.extend(_replace_in_table(t, replacements))
    return replaced_all


def replace_text_in_document(doc: Document, replacements: Dict[str, str]) -> Dict[str, Any]:
    """
    Reemplaza placeholders en:
    - párrafos
    - tablas (y tablas anidadas)

    Retorna un resumen:
    - replaced_count: cantidad de reemplazos detectados
    - replaced_keys: lista única de placeholders reemplazados
    - missing_keys: placeholders presentes en el doc que NO se pudieron reemplazar
    """
    replaced_keys: List[str] = []

    # Reemplazo en párrafos
    for p in doc.paragraphs:
        replaced_keys.extend(_replace_in_paragraph(p, replacements))

    # Reemplazo en tablas
    for t in doc.tables:
        replaced_keys.extend(_replace_in_table(t, replacements))

    # Buscar placeholders que quedaron sin reemplazar
    remaining_placeholders: List[str] = []
    # Escaneamos texto plano del documento (párrafos + celdas)
    # Nota: python-docx no da un "texto completo" único, así que reunimos.
    texts: List[str] = []
    texts.extend([p.text for p in doc.paragraphs])
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.extend([p.text for p in cell.paragraphs])

    for txt in texts:
        for m in PLACEHOLDER_RE.finditer(txt):
            remaining_placeholders.append("${" + m.group(1) + "}")

    replaced_unique = sorted(set(replaced_keys))
    remaining_unique = sorted(set(remaining_placeholders))

    return {
        "replaced_count": len(replaced_keys),
        "replaced_keys": replaced_unique,
        "missing_keys": remaining_unique,
    }


# -----------------------------------------------------------------------------
# Email (SMTP) opcional
# -----------------------------------------------------------------------------

def send_email_with_attachment(
    to_email: str,
    subject: str,
    body: str,
    attachment_path: str,
    attachment_filename: Optional[str] = None,
) -> None:
    """
    Envía un email por SMTP con un archivo adjunto.
    Requiere SMTP_ENABLED=1 y credenciales en env vars.
    """
    if not SMTP_ENABLED:
        raise RuntimeError("SMTP no está habilitado (SMTP_ENABLED=0).")

    import smtplib
    from email.message import EmailMessage

    if not to_email or "@" not in to_email:
        raise ValueError("Email destino inválido.")
    if not SMTP_USER or not SMTP_PASS:
        raise RuntimeError("Faltan SMTP_USER / SMTP_PASS en variables de entorno.")

    attachment_filename = attachment_filename or os.path.basename(attachment_path)
    attachment_filename = sanitize_filename(attachment_filename)

    # Mime type docx
    ctype, encoding = mimetypes.guess_type(attachment_filename)
    if not ctype:
        ctype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    maintype, subtype = ctype.split("/", 1)

    msg = EmailMessage()
    msg["Subject"] = subject or "Documento generado"
    msg["From"] = SMTP_FROM if SMTP_FROM else SMTP_USER
    msg["To"] = to_email
    msg.set_content(body or "Adjunto encontrarás el documento generado.")

    with open(attachment_path, "rb") as f:
        data = f.read()
        msg.add_attachment(data, maintype=maintype, subtype=subtype, filename=attachment_filename)

    if SMTP_DEBUG:
        logger.info("SMTP: host=%s port=%s ssl=%s tls=%s user=%s", SMTP_HOST, SMTP_PORT, SMTP_SSL, SMTP_TLS, SMTP_USER)

    if SMTP_SSL:
        with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT) as smtp:
            smtp.login(SMTP_USER, SMTP_PASS)
            smtp.send_message(msg)
    else:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as smtp:
            if SMTP_TLS:
                smtp.starttls()
            smtp.login(SMTP_USER, SMTP_PASS)
            smtp.send_message(msg)


# -----------------------------------------------------------------------------
# CORS simple
# -----------------------------------------------------------------------------

@app.after_request
def add_cors_headers(resp: Response) -> Response:
    if CORS_ENABLED:
        resp.headers["Access-Control-Allow-Origin"] = CORS_ALLOW_ORIGIN
        resp.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    return resp


@app.route("/<path:_any>", methods=["OPTIONS"])
@app.route("/login.html", methods=["GET"])
def login_html():
    return send_from_directory("static", "login.html")


@app.route("/index.html", methods=["GET"])
def index_html():
    return send_from_directory("static", "index.html")


@app.route("/contacto.html", methods=["GET"])
def contacto_html():
    return send_from_directory("static", "contacto.html")


@app.route("/admin.html", methods=["GET"])
def admin_html():
    return send_from_directory("static", "admin.html")


@app.route("/", methods=["GET"])
def root_html():
    return send_from_directory("static", "index.html")

@app.route("/", methods=["OPTIONS"])
def options_preflight(_any: str = ""):
    return ("", 204)


# -----------------------------------------------------------------------------
# Endpoints
# -----------------------------------------------------------------------------

@app.route("/", methods=["GET"])
def root():
    """
    Opción A (simple): si tenés static/index.html, lo sirve.
    Si no, devuelve un JSON con info.
    """
    index_path = os.path.join(app.static_folder or "static", "index.html")
    if os.path.isfile(index_path):
        return send_file(index_path)
    return jsonify(
        {
            "service": "zeon",
            "status": "ok",
            "endpoints": [
                "GET /health",
                "GET /templates",
                "POST /generate",
                "GET /download/<filename>",
            ],
        }
    )


@app.route("/health", methods=["GET"])
def health():
    ensure_dirs()
    cleanup_outputs()
    return jsonify(
        {
            "ok": True,
            "service": "zeon",
            "template_root": TEMPLATE_ROOT,
            "output_root": OUTPUT_ROOT,
            "smtp_enabled": SMTP_ENABLED,
        }
    )


@app.route("/templates", methods=["GET"])
def templates():
    """
    Lista tipos de plantilla disponibles (carpetas dentro de TEMPLATE_ROOT).
    """
    try:
        types = list_template_types()
        return jsonify({"ok": True, "templates": types})
    except Exception as e:
        logger.exception("Error listando plantillas")
        return jsonify({"ok": False, "error": str(e)}), 500


@dataclass
class GenerateRequest:
    template_type: str
    data: Dict[str, Any]
    output_prefix: str
    email_to: Optional[str] = None
    email_subject: Optional[str] = None
    email_body: Optional[str] = None


def parse_generate_request(req) -> GenerateRequest:
    """
    Acepta:
    - JSON (Content-Type: application/json)
    - form-data / x-www-form-urlencoded (por compatibilidad)

    Formato esperado:
    {
      "template_type": "reparacion_por_fuego",
      "data": { "NOMBRE": "Juan", "DNI": "123" , ... },
      "output_prefix": "Contrato_Juan" (opcional),
      "email": {
         "to": "cliente@correo.com",
         "subject": "Tu documento",
         "body": "Adjunto..."
      } (opcional)
    }
    """
    if req.is_json:
        payload = req.get_json(silent=True) or {}
    else:
        # fallback: si mandan form fields
        payload = dict(req.form or {})
        # Si mandan data como string JSON en form-data:
        if "data" in payload and isinstance(payload["data"], str):
            try:
                payload["data"] = json.loads(payload["data"])
            except Exception:
                # dejamos como string y se convertirá a placeholder único si fuera necesario
                pass

    template_type = (payload.get("template_type") or payload.get("tipo") or payload.get("template") or "").strip()
    data = payload.get("data") or payload.get("datos") or {}
    output_prefix = (payload.get("output_prefix") or payload.get("prefix") or "documento").strip()

    email_to = None
    email_subject = None
    email_body = None
    email_block = payload.get("email") or {}
    if isinstance(email_block, dict):
        email_to = (email_block.get("to") or email_block.get("destino") or "").strip() or None
        email_subject = (email_block.get("subject") or "").strip() or None
        email_body = (email_block.get("body") or "").strip() or None

    # Validaciones mínimas
    if not template_type:
        raise ValueError("Falta 'template_type' (o 'tipo').")
    if not isinstance(data, dict):
        raise ValueError("'data' debe ser un objeto JSON/dict.")

    return GenerateRequest(
        template_type=template_type,
        data=data,
        output_prefix=output_prefix or "documento",
        email_to=email_to,
        email_subject=email_subject,
        email_body=email_body,
    )


@app.route("/generate", methods=["POST"])
def generate():
    """
    Genera un .docx desde una plantilla.

    Respuesta:
    {
      "ok": true,
      "filename": "...docx",
      "download_url": "...",
      "summary": { "replaced_count": 10, "replaced_keys": [...], "missing_keys": [...] },
      "email": { "sent": true/false, "to": "...", "error": "..." }
    }
    """
    ensure_dirs()
    cleanup_outputs()

    try:
        gen_req = parse_generate_request(request)

        template_path = find_template_docx(gen_req.template_type)
        doc = Document(template_path)

        replacements = normalize_replacements(gen_req.data)
        summary = replace_text_in_document(doc, replacements)

        out_name = build_output_filename(prefix=gen_req.output_prefix, ext=".docx")
        out_path = output_path_for(out_name)
        doc.save(out_path)

        # Construir URL de descarga
        if APP_BASE_URL:
            download_url = f"{APP_BASE_URL}/download/{out_name}"
        else:
            # relativa (funciona si llamás desde el mismo host)
            download_url = f"/download/{out_name}"

        email_result = {"sent": False, "to": gen_req.email_to, "error": None}
        if gen_req.email_to:
            try:
                send_email_with_attachment(
                    to_email=gen_req.email_to,
                    subject=gen_req.email_subject or "Documento generado",
                    body=gen_req.email_body or "Adjunto encontrarás el documento generado.",
                    attachment_path=out_path,
                    attachment_filename=out_name,
                )
                email_result["sent"] = True
            except Exception as e:
                logger.exception("Error enviando email")
                email_result["error"] = str(e)

        return jsonify(
            {
                "ok": True,
                "filename": out_name,
                "download_url": download_url,
                "summary": summary,
                "email": email_result,
            }
        )

    except FileNotFoundError as e:
        return jsonify({"ok": False, "error": str(e)}), 404
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    except Exception:
        logger.exception("Error inesperado en /generate")
        return jsonify({"ok": False, "error": "Error interno generando documento."}), 500


@app.route("/download/<filename>")
def download(filename):
    base_dir = os.path.abspath(os.path.dirname(__file__))
    zip_dir = os.path.join(base_dir, "outputs", "zips")

    file_path = os.path.join(zip_dir, filename)

    if not os.path.isfile(file_path):
        abort(404, description="Archivo no encontrado.")

    return send_from_directory(
        directory=zip_dir,
        path=filename,
        as_attachment=True
    )

@app.route("/debug/template/<template_type>", methods=["GET"])
def debug_template(template_type: str):
    """
    Endpoint útil para debug:
    - Te dice qué plantilla .docx se seleccionaría para ese tipo
    - No expone el contenido, solo path relativo y nombre
    """
    try:
        p = find_template_docx(template_type)
        rel = os.path.relpath(p, os.getcwd())
        return jsonify({"ok": True, "template_type": template_type, "template_path": rel})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 400


# -----------------------------------------------------------------------------
# Utilidad opcional: endpoint para copiar un ejemplo de estructura (dev)
# -----------------------------------------------------------------------------

@app.route("/dev/bootstrap", methods=["POST"])
def dev_bootstrap():
    """
    Crea carpetas y un ejemplo mínimo (solo para desarrollo local).
    En Render no es muy útil (disco efímero), pero ayuda localmente.

    Cuerpo:
    {
      "template_type": "ejemplo",
      "copy_from": "plantillas/otra/plantilla.docx" (opcional)
    }
    """
    if os.getenv("FLASK_ENV", "").lower() == "production":
        return jsonify({"ok": False, "error": "Deshabilitado en producción."}), 403

    payload = request.get_json(silent=True) or {}
    template_type = sanitize_filename(payload.get("template_type", "ejemplo"))
    copy_from = payload.get("copy_from", "")

    folder = safe_join(TEMPLATE_ROOT, template_type)
    os.makedirs(folder, exist_ok=True)

    target = os.path.join(folder, "plantilla.docx")
    if copy_from:
        try:
            shutil.copyfile(copy_from, target)
        except Exception as e:
            return jsonify({"ok": False, "error": f"No se pudo copiar: {e}"}), 400

    return jsonify({"ok": True, "created": folder, "template": target})


# -----------------------------------------------------------------------------
# Main (solo para correr local con python app.py)
# -----------------------------------------------------------------------------

if __name__ == "__main__":
    ensure_dirs()
    cleanup_outputs()
    port = _env_int("PORT", 5000)  # Render inyecta PORT
    debug = os.getenv("FLASK_ENV", "").lower() != "production"
    logger.info("Iniciando Zeon en http://0.0.0.0:%s (debug=%s)", port, debug)
    app.run(host="0.0.0.0", port=port, debug=False)
